from docx import Document
import os
import datetime
import subprocess
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

APP_ROOT = os.path.dirname(os.path.abspath(__file__))
RAW_DOC_FOLDER = os.path.join(APP_ROOT, 'static/raw_doc')
FINAL_DOC_FOLDER = os.path.join(APP_ROOT, 'static/documents')


class DocumentManager:

    def generateDocFile(self, issue_obj):
        
        # Load raw file as input 
        
        document = Document(os.path.join(RAW_DOC_FOLDER, 'report.docx'))
        # document_sections = document.sections
        paragraph = document.paragraphs

        # Edit document paragraph by paragraph : 
        desc_text = 'Description : ' + str(issue_obj['description']) 
        paragraph[6].text = desc_text # Description
        
        location_text = 'Location of Damage : ' + str(issue_obj['locality'])
        paragraph[8].text = location_text

        location_type = 'Type of Damage : ' + str(issue_obj['label'])
        paragraph[9].text = location_type

        date_reported = 'Date Reported : ' + str(issue_obj['time'])
        paragraph[10].text = date_reported

        priority_text = 'Priority : Severe'
        paragraph[11].text = priority_text

        reported_by = 'Reported By : ' + str(issue_obj['uploaded_by'])
        paragraph[12].text = reported_by 

        current_status = 'Current Status : ' + str("Not fixed")
        paragraph[13].text = current_status

        timestamp = str(datetime.datetime.now().strftime('%H:%M:%S'))

        filename = issue_obj['category'] + '_' + issue_obj['issue_id'] + '_report_' + timestamp + '.docx'
        document.save(os.path.join(FINAL_DOC_FOLDER, filename))

        final_path = os.path.join(FINAL_DOC_FOLDER, filename)
        self.doc2pdf(final_path)

        return 'Done'


    def move_table_after(self, table, paragraph):
        tbl, p = table._tbl, paragraph._p
        p.addnext(tbl)

    def move_image_after(self, image, paragraph):
        img, p = image._inline, paragraph._p
        p.addnext(img)


    def generateDocTable(self, issue_obj):

        # Load raw file as input 
        
        document = Document(os.path.join(RAW_DOC_FOLDER, 'report_table.docx'))
        # document_sections = document.sections
        paragraph = document.paragraphs

        font_normal = document.styles['Normal'].font
        # font.name = 'Proxima Nova Rg'
        font_normal.name = 'Rubik'

        font_heading_1 = document.styles['Title'].font
        # font.name = 'Proxima Nova Rg'
        font_heading_1.name = 'Rubik'

        # Edit document paragraph by paragraph : 
        desc_text = 'Description : ' + str(issue_obj['description']) 
        paragraph[6].text = desc_text # Description


        details_preceding_paragraph = paragraph[8]  # however you get this paragraph
        table = document.add_table(rows=6, cols=2)
        table.style = None

        # 1st row
        cell_0_0 = table.cell(0, 0)
        cell_0_0.text = 'Location'

        cell_0_1 = table.cell(0, 1)
        cell_0_1.text = str(issue_obj['locality'])

        # 2nd row
        cell_1_0 = table.cell(1, 0)
        cell_1_0.text = 'Type of Damage'

        cell_1_1 = table.cell(1, 1)
        cell_1_1.text = str(issue_obj['label'])

        # 3rd row
        cell_2_0 = table.cell(2, 0)
        cell_2_0.text = 'Date Reported'

        cell_2_1 = table.cell(2, 1)
        cell_2_1.text = str(issue_obj['time'])

        # 4th row
        cell_3_0 = table.cell(3, 0)
        cell_3_0.text = 'Priority'

        cell_3_1 = table.cell(3, 1)
        cell_3_1.text = str('Severe')

        # 5th row
        cell_4_0 = table.cell(4, 0)
        cell_4_0.text = 'Reported by'

        cell_4_1 = table.cell(4, 1)
        cell_4_1.text = str(issue_obj['uploaded_by'])

        # 6th row
        cell_5_0 = table.cell(5, 0)
        cell_5_0.text = 'Current Status'

        cell_5_1 = table.cell(5, 1)
        cell_5_1.text = str('Not Fixed')

        self.move_table_after(table, details_preceding_paragraph)

        image_1_path = str(issue_obj['image_one_full'])
        run = paragraph[11].add_run()
        paragraph[11].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run.add_picture(image_1_path, width=Inches(3.0))

        image_2_path = str(issue_obj['image_two_full'])
        run = paragraph[13].add_run()
        paragraph[13].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run.add_picture(image_2_path, width=Inches(3.0))

        image_3_path = str(issue_obj['detection'])
        run = paragraph[15].add_run()
        paragraph[15].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run.add_picture(image_3_path, width=Inches(3.0))
        
        timestamp = str(datetime.datetime.now().strftime('%H:%M:%S'))

        filename = issue_obj['category'] + '_' + issue_obj['issue_id'] + '_report_table_' + timestamp + '.docx'
        document.save(os.path.join(FINAL_DOC_FOLDER, filename))

        final_path = os.path.join(FINAL_DOC_FOLDER, filename)
        self.doc2pdf(final_path)

        pdf_name = filename.replace('docx', 'pdf')
        doc_uri = '/documents/%s' % (pdf_name)
        return doc_uri


    def doc2pdf(self, file_path):
        args = ['libreoffice', '--headless', '--convert-to', 'pdf', '--outdir', FINAL_DOC_FOLDER, file_path]
        process = subprocess.run(args, stdout=subprocess.PIPE, stderr=subprocess.PIPE, timeout=10)
        return 'done'
  